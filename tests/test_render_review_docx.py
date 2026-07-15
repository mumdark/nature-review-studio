"""Re-render EN + ZH payloads via render_review_docx.py and assert:
  * non-empty valid docx
  * paragraphs >= 30, tables >= 1
  * Heading 1/2/3/4 + Normal styles present
  * companion .md emitted (locked V1.4 dual-output contract)

Self-contained: builds two minimal payloads in tmp_path. No external
fixtures required. Resolves the renderer via NRS_ROOT (or legacy NRL_ROOT).

Run with:
    python tests/test_render_review_docx.py
"""
from __future__ import annotations
import json
import os
import subprocess
import sys
import tempfile
from pathlib import Path

HERE = Path(__file__).resolve().parent
ENV_ROOT = os.environ.get("NRS_ROOT") or os.environ.get("NRL_ROOT")
NRS_ROOT = Path(ENV_ROOT).resolve() if ENV_ROOT else HERE.parents[1]
sys.path.insert(0, str(NRS_ROOT / "scripts" / "python_docx"))
PYTHON = sys.executable
SCRIPT = str(NRS_ROOT / "scripts" / "render_review_docx.py")

REQUIRED_STYLES = {"Heading 1", "Heading 2", "Heading 3", "Heading 4", "Normal"}


def _minimal_payload(lang: str, case_id: str) -> dict:
    if lang == "zh":
        return {
            "entry": "review",
            "case_id": case_id,
            "language": "zh",
            "editor_letter": {
                "revision_level": "Major revision",
                "paragraphs": [
                    "感谢投稿。请逐条回应各位审稿人的意见后重新提交。",
                    "本稿件需进行较大修改。",
                ],
            },
            "reviewers": [
                {
                    "role_label": "机制审稿人",
                    "overall_paragraph": "整体而言,论文选题有意义,但实验证据需要进一步强化。",
                    "major_concerns": [
                        {"id": "M1", "title": "机制证据", "body": "请补充rescue实验。"},
                        {"id": "M2", "title": "Power analysis", "body": "请提供power analysis。"},
                    ],
                    "minor_concerns": [
                        {"id": "m1", "title": "图轴", "body": "请将图3纵轴改为0-100%。"},
                    ],
                },
                {
                    "role_label": "统计审稿人",
                    "overall_paragraph": "统计与可复现性方面需要补强。",
                    "major_concerns": [
                        {"id": "M3", "title": "稳健性", "body": "请补充bootstrap稳健性分析。"},
                        {"id": "M4", "title": "代码数据", "body": "请上传分析代码与匿名数据。"},
                    ],
                    "minor_concerns": [
                        {"id": "m2", "title": "图注", "body": "请清晰图注。"},
                    ],
                },
            ],
            "consensus": [
                {"id": "C1", "description": "机制证据需补强", "raised_by": ["审稿人 1", "审稿人 2"]},
                {"id": "C2", "description": "统计需补强", "raised_by": ["审稿人 2"]},
            ],
            "revision_task_table": {
                "headers": ["编号", "审稿人", "问题", "策略", "状态"],
                "rows": [
                    ["T1", "审稿人 1", "机制证据", "add_experiment", "TODO_EXPERIMENT"],
                    ["T2", "审稿人 2", "稳健性", "add_robustness_analysis", "TODO_ANALYSIS"],
                ],
            },
        }
    return {
        "entry": "review",
        "case_id": case_id,
        "language": "en",
        "editor_letter": {
            "revision_level": "Major revision",
            "paragraphs": [
                "Thank you for submitting your manuscript.",
                "A major revision is required.",
            ],
        },
        "reviewers": [
            {
                "role_label": "Mechanism Reviewer",
                "overall_paragraph": "Overall, the topic is meaningful, but experimental evidence needs to be strengthened.",
                "major_concerns": [
                    {"id": "M1", "title": "Mechanism evidence", "body": "Please add a rescue experiment."},
                    {"id": "M2", "title": "Power analysis", "body": "Please provide a power analysis."},
                ],
                "minor_concerns": [
                    {"id": "m1", "title": "Figure axis", "body": "Please make Figure 3 axis 0-100%."},
                ],
            },
            {
                "role_label": "Statistics Reviewer",
                "overall_paragraph": "Statistical rigor and reproducibility need reinforcement.",
                "major_concerns": [
                    {"id": "M3", "title": "Robustness", "body": "Please add bootstrap robustness analysis."},
                    {"id": "M4", "title": "Code/data sharing", "body": "Please upload analysis code and anonymized data."},
                ],
                "minor_concerns": [
                    {"id": "m2", "title": "Caption clarity", "body": "Please clarify the caption."},
                ],
            },
        ],
        "consensus": [
            {"id": "C1", "description": "Mechanism evidence needs strengthening", "raised_by": ["Reviewer 1", "Reviewer 2"]},
            {"id": "C2", "description": "Statistical and reproducibility concerns", "raised_by": ["Reviewer 2"]},
        ],
        "revision_task_table": {
            "headers": ["ID", "Reviewer", "Concern", "Strategy", "Status"],
            "rows": [
                ["T1", "Reviewer 1", "Mechanism evidence", "add_experiment", "TODO_EXPERIMENT"],
                ["T2", "Reviewer 2", "Robustness", "add_robustness_analysis", "TODO_ANALYSIS"],
            ],
        },
    }


def main() -> None:
    ok = True
    if not Path(SCRIPT).exists():
        print(f"FAIL: renderer not found at {SCRIPT}")
        sys.exit(2)
    with tempfile.TemporaryDirectory() as tmp:
        tmpdir = Path(tmp)
        for lang, case_id in (("en", "release_en"), ("zh", "release_zh")):
            payload_path = tmpdir / f"payload_{lang}.json"
            out_docx = tmpdir / f"review_{lang}.docx"
            payload_path.write_text(json.dumps(_minimal_payload(lang, case_id), ensure_ascii=False, indent=2), encoding="utf-8")
            r = subprocess.run(
                [PYTHON, SCRIPT, "--payload", str(payload_path), "--out", str(out_docx)],
                capture_output=True, text=True,
            )
            if r.returncode != 0:
                print(f"FAIL render {out_docx.name}: rc={r.returncode} stderr={r.stderr.strip()}")
                ok = False
                continue
            from docx import Document
            d = Document(str(out_docx))
            styles = {p.style.name for p in d.paragraphs}
            missing = REQUIRED_STYLES - styles
            if missing:
                print(f"FAIL {out_docx.name}: missing styles {missing}")
                ok = False
                continue
            if len(d.paragraphs) < 30 or len(d.tables) < 1:
                print(f"FAIL {out_docx.name}: too small (paragraphs={len(d.paragraphs)} tables={len(d.tables)})")
                ok = False
                continue
            md_path = out_docx.with_suffix(".md")
            if not md_path.exists() or md_path.stat().st_size == 0:
                print(f"FAIL {out_docx.name}: missing companion .md")
                ok = False
                continue
            print(f"PASS review_{lang}.docx: paragraphs={len(d.paragraphs)} tables={len(d.tables)} styles ok md={md_path.stat().st_size}B")
    if not ok:
        sys.exit(1)


if __name__ == "__main__":
    main()
